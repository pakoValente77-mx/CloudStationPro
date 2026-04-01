#!/usr/bin/env python3
"""Genera documento Word con solicitud de permisos de red para producción."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_bg(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear'
    })
    shading.append(elm)


def add_styled_table(doc, data, header_color, alt_color):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for j, text in enumerate(data[0]):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_bg(cell, header_color)

    # Data rows
    for i in range(1, len(data)):
        for j, text in enumerate(data[i]):
            cell = table.rows[i].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(text)
            run.font.size = Pt(7.5)
            if i % 2 == 0:
                set_cell_bg(cell, alt_color)

    return table


def main():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('COMISIÓN FEDERAL DE ELECTRICIDAD')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 100, 0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Subgerencia Regional de Generación Hidroeléctrica Grijalva')
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Solicitud de Permisos de Red para Servidor de Producción')
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Sistema CloudStation — Monitoreo Hidrológico GOES')
    run.font.size = Pt(11)
    run.font.italic = True

    doc.add_paragraph()

    doc.add_paragraph('Fecha: 30 de marzo de 2026')
    doc.add_paragraph('Servidor de producción: 192.168.1.72 (Windows Server)')
    doc.add_paragraph('Responsable: Subgerencia Regional de Generación Hidroeléctrica Grijalva')

    doc.add_paragraph()

    # ── Sección 1: Conexiones externas ──
    doc.add_heading('1. Conexiones Externas (Salida a Internet)', level=1)

    doc.add_paragraph(
        'Las siguientes conexiones son necesarias para que el sistema CloudStation '
        'pueda recibir datos satelitales GOES, pronósticos meteorológicos y enviar '
        'notificaciones por correo electrónico.'
    )

    ext_data = [
        ['#', 'Host / URL', 'Puerto', 'Protocolo', 'Servicio',
         'Componente', 'Justificación'],
        ['1', 'lrgseddn1.cr.usgs.gov', '16003', 'TCP (DDS)',
         'LRGS/DDS\n(USGS)', 'mycloud_all_timescale.py\nrecover_goes.py',
         'Servidor primario de datos satelitales GOES del USGS. Transmite mediciones hidrológicas (nivel, precipitación) de 135 estaciones en tiempo real.'],
        ['2', 'lrgseddn2.cr.usgs.gov', '16003', 'TCP (DDS)',
         'LRGS/DDS\n(USGS)', 'mycloud_all_timescale.py\nrecover_goes.py',
         'Servidor secundario (failover) de datos GOES.'],
        ['3', 'lrgseddn3.cr.usgs.gov', '16003', 'TCP (DDS)',
         'LRGS/DDS\n(USGS)', 'mycloud_all_timescale.py\nrecover_goes.py',
         'Servidor terciario (failover) de datos GOES.'],
        ['4', 'atlas16.ddns.net', '5432', 'TCP',
         'PostgreSQL\n(Desarrollo)', 'mycloud_all_timescale.py\ntransfer_datos.py',
         'Base de datos de desarrollo/respaldo. Transferencia de datos y recuperación ante pérdida de datos en producción.'],
        ['5', 'atlas16.ddns.net', '1433', 'TCP (TDS)',
         'SQL Server\n(Desarrollo)', 'mycloud_all_timescale.py',
         'Base de datos SQL Server de desarrollo. Catálogo de estaciones y configuración del sistema.'],
        ['6', '200.4.8.36', '21', 'TCP (FTP)',
         'FTP\n(Pronóstico)', 'sync_rain_forecast.py',
         'Servidor FTP para descarga de archivos CSV de pronóstico de lluvia. Actualización cada 6 horas.'],
        ['7', '159.16.6.20', '25', 'TCP (SMTP)',
         'Correo\nElectrónico', 'CloudStationWeb\n(NotificationService)',
         'Servidor SMTP institucional para envío de alertas de nivel crítico, reportes automáticos y notificaciones del sistema de alerta temprana.'],
        ['8', 'api.open-meteo.com', '443', 'HTTPS',
         'API Meteorológica\n(Open-Meteo)', 'CloudStationWeb\n(WeatherService)',
         'API pública de pronóstico meteorológico. Condiciones actuales y pronóstico para las presas monitoreadas.'],
        ['9', 'grijalva.cfe.mx', '8010', 'HTTP',
         'Eureka\n(Service Discovery)', 'grijalva-hydro-model-service\ngrijalva-rain-forecast-service',
         'Servidor de registro de microservicios (Spring Eureka). Los servicios Java se registran aquí.'],
        ['10', 'grijalva.cfe.mx\ndev.grijalva.cfe.mx', '5432', 'TCP',
         'PostgreSQL\n(Microservicios)', 'grijalva-hydro-model-service\ngrijalva-rain-forecast-service',
         'Bases de datos PostgreSQL de los microservicios Java de simulación hidrológica y pronóstico de lluvia.'],
        ['11', 'mesonet.agron.iastate.edu', '443', 'HTTPS',
         'WMS GOES-16\n(Iowa State / NOAA)', 'CloudStationWeb\n(RainForecast map)',
         'Imágenes satelitales GOES-16 en tiempo real vía WMS. Capas: infrarrojo (ch13, 24/7) y visible (ch02, diurno). Usadas en el mapa de pronóstico de lluvia.'],
        ['12', '{s}.google.com', '80 / 443', 'HTTP/HTTPS',
         'Google Maps\n(Tiles)', 'CloudStationWeb\n(Map, RainForecast,\nStationAdmin)',
         'Mosaicos base del mapa: terreno (lyrs=p) y satélite (lyrs=s,h). Se usan en todas las vistas con mapa Leaflet.'],
        ['13', 'api.rainviewer.com\ntilecache.rainviewer.com', '443', 'HTTPS',
         'RainViewer\n(Radar precipitación)', 'CloudStationWeb\n(RainForecast map)',
         'API y tiles de radar de precipitación en tiempo real. Muestra animación de lluvia sobre el mapa de la cuenca Grijalva.'],
        ['14', 'cdn.jsdelivr.net', '443', 'HTTPS',
         'CDN Librerías JS\n(jsdelivr)', 'CloudStationWeb\nmonitor.html',
         'CDN para librerías JavaScript: jQuery, Fomantic-UI, Chart.js, Three.js, Leaflet. Necesario para que la interfaz web cargue correctamente.'],
        ['15', 'fcm.googleapis.com', '443', 'HTTPS',
         'Firebase Cloud\nMessaging (FCM)', 'CloudStationWeb\n(PushNotificationService)',
         'Push notifications a dispositivos móviles y web. Envía alertas de nivel crítico en presas, notificaciones de chat y avisos operativos a la app móvil.'],
        ['16', 'oauth2.googleapis.com\nfcmregistrations.googleapis.com', '443', 'HTTPS',
         'Firebase Auth\n/ Registration', 'CloudStationWeb\n(PushNotificationService)',
         'Autenticación OAuth2 del servidor con Firebase y registro de tokens de dispositivos. Requerido para que FCM funcione.'],
        ['17', 'firebaseinstallations.googleapis.com', '443', 'HTTPS',
         'Firebase\nInstallations', 'App móvil\n(cliente)',
         'Registro e instalación de la app móvil en Firebase. Los dispositivos contactan este servicio para obtener su token FCM.'],
        ['18', 'smtp.gmail.com', '587', 'TCP (SMTP\nSTARTTLS)',
         'Gmail SMTP\n(Desarrollo)', 'grijalva-notification-service\n(Spring Boot)',
         'Servidor SMTP alternativo utilizado por el microservicio Java de notificaciones en ambiente de desarrollo. Requiere TLS.'],
    ]

    add_styled_table(doc, ext_data, '003366', 'E8F0FE')
    doc.add_paragraph()

    # ── Sección 2: Conexiones locales ──
    doc.add_heading('2. Conexiones Locales (Red Interna 192.168.1.x)', level=1)

    doc.add_paragraph(
        'Las siguientes conexiones operan dentro del servidor de producción o en la red local. '
        'Se incluyen para referencia del administrador de red.'
    )

    local_data = [
        ['#', 'Host / URL', 'Puerto', 'Protocolo', 'Servicio',
         'Componente', 'Justificación'],
        ['1', 'localhost (127.0.0.1)', '5432', 'TCP',
         'PostgreSQL\n(TimescaleDB)', 'mycloud_all_timescale.py\nCloudStationWeb',
         'Base de datos de series de tiempo local. Almacena mediciones GOES, resúmenes horarios/diarios y pronósticos.'],
        ['2', 'localhost (127.0.0.1)', '1433', 'TCP (TDS)',
         'SQL Server\n(IGSCLOUD_PRO)', 'CloudStationWeb\nmycloud_all_timescale.py',
         'Base de datos relacional local. Catálogo de estaciones, usuarios, permisos, documentos y configuración.'],
        ['3', 'localhost (127.0.0.1)', '5555', 'HTTP',
         'Flask Monitor', 'monitor.html\n(navegador)',
         'Micro-servidor web de monitoreo interno. Estado del colector GOES, últimas recepciones y estadísticas.'],
        ['4', '192.168.1.72', '80 / 443', 'HTTP/HTTPS',
         'IIS\n(ASP.NET Core)', 'Navegador web\n(usuarios finales)',
         'Aplicación web CloudStationWeb publicada en IIS. Interfaz principal del sistema de monitoreo hidrológico. Puerto 80 HTTP, puerto 443 HTTPS (si se configura certificado SSL).'],
        ['5', '192.168.1.72', '80/hubs/chat\n443/hubs/chat', 'WebSocket\n(sobre HTTP/HTTPS)',
         'SignalR Hub\n(Chat tiempo real)', 'CloudStationWeb\n(navegador / app)',
         'Hub de chat en tiempo real usando ASP.NET SignalR. Opera sobre el mismo puerto 80/443 de IIS. No requiere puerto adicional, solo que el firewall permita conexiones WebSocket.'],
        ['6', 'grijalva-notification\n-service (Eureka)', 'Dinámico', 'HTTP',
         'Microservicio\nNotificaciones\n(Spring Boot)', 'grijalva-hydro-model-service\nCloudStationWeb',
         'Microservicio Java que envía correos electrónicos vía REST (POST /v1/mail/send). Se descubre dinámicamente a través de Eureka.'],
    ]

    add_styled_table(doc, local_data, '2E7D32', 'E8F5E9')
    doc.add_paragraph()

    # ── Sección 3: Resumen de puertos ──
    doc.add_heading('3. Resumen de Puertos a Abrir (Firewall de Salida)', level=1)

    summary_data = [
        ['Puerto', 'Protocolo', 'Destinos', 'Dirección'],
        ['16003', 'TCP', 'lrgseddn1/2/3.cr.usgs.gov', 'SALIDA'],
        ['5432', 'TCP', 'atlas16.ddns.net, grijalva.cfe.mx, dev.grijalva.cfe.mx', 'SALIDA'],
        ['1433', 'TCP', 'atlas16.ddns.net', 'SALIDA'],
        ['21', 'TCP (FTP)', '200.4.8.36', 'SALIDA'],
        ['25', 'TCP (SMTP)', '159.16.6.20', 'SALIDA'],
        ['443', 'HTTPS', 'api.open-meteo.com, mesonet.agron.iastate.edu,\napi.rainviewer.com, tilecache.rainviewer.com,\ncdn.jsdelivr.net, fcm.googleapis.com,\noauth2.googleapis.com,\nfcmregistrations.googleapis.com,\nfirebaseinstallations.googleapis.com', 'SALIDA'],
        ['587', 'TCP (SMTP TLS)', 'smtp.gmail.com (desarrollo)', 'SALIDA'],
        ['80', 'HTTP', '{s}.google.com (tiles de mapas)', 'SALIDA'],
        ['8010', 'HTTP', 'grijalva.cfe.mx', 'SALIDA'],
    ]

    add_styled_table(doc, summary_data, '8B0000', 'FFF0F0')
    doc.add_paragraph()

    # ── Sección 4: Notas ──
    doc.add_heading('4. Notas Importantes', level=1)

    notas = [
        'Todos los puertos indicados son de SALIDA (el servidor 192.168.1.72 inicia la conexión hacia los destinos externos).',
        'Los servidores LRGS del USGS (puerto 16003) son CRÍTICOS. Sin acceso a estos, el sistema no puede recibir datos satelitales GOES de las 135 estaciones hidrológicas monitoreadas.',
        'El servidor SMTP (puerto 25) es necesario para el sistema de alerta temprana que notifica niveles críticos en presas.',
        'El FTP (puerto 21) se puede configurar en modo pasivo si el firewall lo requiere (puertos dinámicos 1024-65535).',
        'La API de Open-Meteo (puerto 443) es una API REST pública que solo requiere HTTPS saliente.',
        'Las imágenes satelitales GOES-16 de la NOAA se obtienen vía WMS (HTTPS) desde mesonet.agron.iastate.edu (Iowa State University). Sin este acceso no se muestran las capas de satélite en el mapa.',
        'RainViewer (api.rainviewer.com / tilecache.rainviewer.com) provee radar de precipitación en tiempo real. Es una API pública HTTPS.',
        'Google Maps tiles ({s}.google.com, puerto 80/443) proveen los mapas base de terreno y satélite para todas las vistas con mapa.',
        'El CDN jsdelivr.net (puerto 443) sirve librerías JavaScript esenciales (jQuery, Leaflet, Chart.js, etc.). Sin acceso, la interfaz web no carga.',
        'Firebase Cloud Messaging (fcm.googleapis.com, puerto 443) es necesario para enviar push notifications a dispositivos móviles y web. Se usa para alertas de nivel crítico, notificaciones de chat y avisos operativos.',
        'El chat en tiempo real usa SignalR (WebSocket) sobre el mismo puerto 80/443 de IIS. No requiere puerto adicional de firewall, pero los navegadores deben poder establecer conexiones WebSocket al servidor.',
        'El microservicio grijalva-notification-service (Spring Boot) se registra en Eureka y envía correos vía POST /v1/mail/send. Se conecta al SMTP institucional (159.16.6.20:25) o a smtp.gmail.com:587 en desarrollo.',
        'Las conexiones a atlas16.ddns.net son temporales durante la fase de migración a producción. Una vez consolidado, solo se necesitarán los servidores LRGS y los servicios internos.',
    ]

    for nota in notas:
        p = doc.add_paragraph(nota, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph()

    # Firmas
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('_' * 45 + '                    ' + '_' * 45)
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(
        'Solicitante                                                         '
        '             Visto Bueno Administrador de Red'
    )
    run.font.size = Pt(10)

    output = '/Users/subgerenciagrijalva/CFE/CloudStation/Solicitud_Permisos_Red_CloudStation.docx'
    doc.save(output)
    print(f'Documento generado: {output}')


if __name__ == '__main__':
    main()
